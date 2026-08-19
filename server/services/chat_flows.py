# _*_ coding : UTF-8 _*_
"""
子流程（Flow）注册表：intent → {system_prompt, tools, data_prep}

Langflow 思想映射：
- 每个大类 = 一条独立子流程（Flow），职责单一、可独立测试；
- 一条子流程由「数据预取节点 data_prep（把该问题需要的数据查好注入 prompt）
  + 专用系统提示词节点 system_prompt（角色 + 工具 + 规则，不给死步骤）
  + 工具子集 tools（只放该流程需要的工具）」组成；
- 路由器输出 intent 后，chat_stream 按 intent 选中一个 Flow 执行；other 走原 RAG 万能链路。

报告类特殊：回答完成后由 chat_stream 调 generate_report_docx 把 Markdown 转 Word 供下载。
"""
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional

from loguru import logger

from core import state
from core.config import config
from core.db import pg_store
from services.agent_tools import (
    _fmt_student_table,
    _fmt_class_table,
    _fmt_top_error,
    expand_class_label,
)

# ============================================================
# 数据预取节点（data_prep）：返回注入 prompt 的上下文文本，查不到返回 ""
# ============================================================
def _fix_name(name: Optional[str]) -> Optional[str]:
    """修正弱姓名抽取的尾部误吞（如"张三数"来自"张三数学"）。
    仅去掉学业/成绩类尾字，尽力而为；查不到时模型会如实说明。"""
    if name and len(name) > 2 and name[-1] in "数学分试测验业绩科目":
        return name[:-1]
    return name


def _class_explain_text(user_msg_class: str, class_names: List[str]) -> str:
    """生成班级粒度说明（前置注入，让模型第一句就解释合并班关系，避免拿错班当主体）。"""
    if not class_names:
        return ""
    merged: List[str] = []
    for c in sorted(set(class_names)):
        ex = expand_class_label(c)
        if ex != c and ex not in merged:
            merged.append(ex)
    if not merged:
        return ""
    return (
        "【重要：班级匹配说明】\n"
        f"老师查询的班级「{user_msg_class}」在数据中没有精确对应的班级名，"
        "数据中的班级是合并班，匹配到的有：" + "；".join(merged) + "。\n"
        "合并班数据未按自然班单独拆分，无法单独给出其中某个自然班（如「软件1班」）的数据。"
        "回答时第一句必须先说明这一匹配情况，并询问老师要看哪个合并班（学生数最多的通常就是最近上传的数据）；"
        "绝对不要把某个合并班的数据直接冒充为「软件1班」，也不得忽略其他匹配的合并班。"
    )



def _prep_student_score(user_msg: str, params: Dict[str, Any]) -> str:
    """成绩-个人：行级明细（每个实验的分数）+ 聚合卡。"""
    sid, name = params.get("student_id"), _fix_name(params.get("name"))
    rows = pg_store.student_scores(student_id=sid, name=name, limit=200)
    if not rows:
        return ""
    parts = ["## 学生个人成绩明细（每个实验一条，按时间排序）"]
    lines = ["| 实验/章节 | 数据源 | 分数 | 薄弱任务数 | 记录时间 |",
             "|---|---|---|---|---|"]
    for r in rows[:40]:
        lines.append(
            f"| {r.get('experiment_name','')} | {r.get('source_type','')} | "
            f"{r.get('final_score','')} | {r.get('weak_count','')} | {r.get('created_at','')} |"
        )
    parts.append("\n".join(lines))
    agg = pg_store.student_summary(student_id=sid, name=name, limit=20)
    if agg:
        parts.append("## 学生聚合（全程平均）\n" + _fmt_student_table(agg))
    return "\n\n".join(parts)


def _prep_class_score(user_msg: str, params: Dict[str, Any]) -> str:
    """成绩-班级：班级粒度说明（前置）+ 班级×实验聚合。"""
    rows = pg_store.class_summary(
        class_name=params.get("class_name"),
        experiment_name=params.get("experiment_name"),
        limit=50,
    )
    if not rows:
        return ""
    parts: List[str] = []
    cls_names = sorted({r.get("class_name", "") for r in rows})
    explain = _class_explain_text(str(params.get("class_name") or ""), cls_names)
    if explain:
        parts.append(explain)
    parts.append("## 班级×实验聚合统计\n" + _fmt_class_table(rows))
    return "\n\n".join(parts)


def _prep_trend(user_msg: str, params: Dict[str, Any]) -> str:
    """成绩-趋势：跨实验/批次分数序列。只有一期数据时要如实说明无法判断趋势。"""
    rows = pg_store.student_trend(
        student_id=params.get("student_id"),
        name=_fix_name(params.get("name")),
        class_name=params.get("class_name"),
        limit=100,
    )
    if not rows:
        return ""
    lines = ["## 成绩分期序列（按时间先后排列，可用于判断进步/退步）"]
    lines.append("| 顺序 | 实验/章节 | 数据源 | 平均分 | 记录条数 | 最早时间 | 最晚时间 |")
    lines.append("|---|---|---|---|---|---|---|")
    for i, r in enumerate(rows[:30], 1):
        lines.append(
            f"| {i} | {r.get('experiment_name','')} | {r.get('source_type','')} | "
            f"{r.get('avg_score','')} | {r.get('record_count','')} | "
            f"{r.get('first_seen','')} | {r.get('last_seen','')} |"
        )
    if len(rows) <= 1:
        lines.append("\n注意：该主体只有一期成绩数据，无法判断趋势变化，请如实告知老师数据不足。")
    return "\n".join(lines)


def _prep_knowledge_gap(user_msg: str, params: Dict[str, Any]) -> str:
    """知识点：班级高频错误/薄弱知识点 + 知识库命中的知识点。"""
    parts: List[str] = []
    ad = state.latest_agg_data or {}
    top_error = ad.get("top_error") or []
    if top_error:
        parts.append(_fmt_top_error(top_error))
    if state.knowledge_base:
        try:
            from analysis.knowledge_builder import search_knowledge
            hits = search_knowledge(user_msg, state.knowledge_base)[:8]
            if hits:
                lines = ["## MOOC 知识库相关知识点（供结合错题解释）"]
                for i, rec in enumerate(hits, 1):
                    name = str(rec.get("视频/知识点名称") or rec.get("name") or "").strip()
                    unit = str(rec.get("MOOC教学单元") or "").strip()
                    if not name:
                        continue
                    tag = "、".join(x for x in [rec.get("知识领域", ""), unit, rec.get("项目名称", "")] if x)
                    lines.append(f"{i}. {name}" + (f"（{tag}）" if tag else ""))
                if len(lines) > 1:
                    parts.append("\n".join(lines))
        except Exception as e:
            logger.warning("knowledge_gap 知识库检索失败: %s", e)
    return "\n\n".join(parts)


def _prep_warning(user_msg: str, params: Dict[str, Any]) -> str:
    """预警：已有预测结论 + 薄弱学生 Top。"""
    parts: List[str] = []
    ad = state.latest_agg_data or {}
    prediction = str(ad.get("prediction_text") or "").strip()
    if prediction:
        parts.append(f"## 已有成绩预警/预测结论\n{prediction}")
    rows = pg_store.student_summary(
        class_name=params.get("class_name"),
        min_weak_rate=30.0,
        limit=50,
        sort_by="weak_rate_percent",
        sort_desc=True,
    )
    if rows:
        parts.append("## 薄弱学生 Top（薄弱率≥30%，来自 v_student_summary）\n" + _fmt_student_table(rows))
    return "\n\n".join(parts)


def _prep_attendance(user_msg: str, params: Dict[str, Any]) -> str:
    """考勤：最近上传的考勤结果（数据在内存 state，未落库；为空时如实说明）。"""
    results = state.latest_attendance_results or []
    if not results:
        return "【当前没有可用的考勤数据】——请先上传考勤 CSV/XLSX 文件，再查询出勤情况。"
    lines = ["## 课堂考勤数据（最近上传）"]
    for a in results[:40]:
        if isinstance(a, dict):
            name = a.get("课堂名称") or a.get("lecture") or a.get("name") or str(a)
            absent = a.get("缺勤人数") or a.get("absent") or ""
            late = a.get("迟到人数") or a.get("late") or ""
            line = f"- {name}"
            if absent:
                line += f"：缺勤 {absent} 人"
            if late:
                line += f"，迟到 {late} 人"
            lines.append(line)
        else:
            lines.append(f"- {a}")
    if len(lines) == 1:
        return ""
    return "\n".join(lines)


def _prep_report(user_msg: str, params: Dict[str, Any]) -> str:
    """报告：班级聚合 + 学生整体 Top + 薄弱学生，作为写报告的素材。"""
    parts: List[str] = []
    rows = pg_store.class_summary(
        class_name=params.get("class_name"),
        experiment_name=params.get("experiment_name"),
        limit=100,
    )
    if rows:
        parts.append("## 班级×实验聚合数据（报告数据来源）\n" + _fmt_class_table(rows))
    stu = pg_store.student_summary(
        class_name=params.get("class_name"),
        limit=100,
        sort_by="weak_rate_percent",
        sort_desc=True,
    )
    if stu:
        parts.append("## 学生按薄弱率排序 Top\n" + _fmt_student_table(stu))
    return "\n\n".join(parts)


# ============================================================
# 子流程注册表：intent → 子流程
# tools 为 AGENT_TOOLS 中的工具名子集；other 不在此表，走原 RAG 万能链路
# ============================================================
FLOWS: Dict[str, Dict[str, Any]] = {
    "score_student": {
        "system_prompt": (
            "你是学生成绩分析助手，面向教师。\n"
            "你有工具 query_student 查学生聚合、query_student_trend 查分期成绩序列；上方数据已按问题预取，优先使用。\n"
            "规则：\n"
            "- 所有数字结论必须来自上方数据或工具返回的真实数据，严禁编造；查不到的学生/实验如实说明缺少数据\n"
            "- 学生信息不完整（没给姓名/学号）时，先说明缺什么，不要瞎猜\n"
            "- 回答用中文，结论放前面，分析放后面"
        ),
        "tools": ["query_student", "query_student_trend"],
        "max_tool_rounds": 3,
        "data_prep": _prep_student_score,
    },
    "score_class": {
        "system_prompt": (
            "你是班级成绩分析助手，面向教师。\n"
            "你有工具 query_class_aggregate 查班级×实验聚合、query_student 查学生聚合；上方数据已按问题预取，优先使用。\n"
            "规则：\n"
            "- 所有数字结论必须来自上方数据或工具返回的真实数据，严禁编造\n"
            "- 班级名可能是合并班（如「软件1-3班」=软件1班+软件2班+软件3班，数据未按自然班拆分）。"
            "当老师查询的班级名（如「软件1班」）在数据中没有精确存在、只匹配到合并班时，"
            "回答第一句必须先说明匹配到的合并班及其包含的自然班、并请老师确认要哪个合并班；"
            "绝对不要把某个合并班的数据直接当作该班数据，也不得忽略其他匹配的合并班\n"
            "- 涉及对比时（班级对比/实验对比），把对比结论明确说清楚，用表格呈现关键指标\n"
            "- 回答用中文，结论放前面，分析放后面"
        ),
        "tools": ["query_class_aggregate", "query_student"],
        "max_tool_rounds": 4,
        "data_prep": _prep_class_score,
    },
    "score_trend": {
        "system_prompt": (
            "你是成绩趋势分析助手，面向教师。\n"
            "你有工具 query_student_trend 查分期序列、query_student 查学生聚合；上方【成绩分期序列】已按时间先后列出各实验/批次平均分。\n"
            "规则：\n"
            "- 判断进步/退步必须基于分期序列前后对比，序列只有一期时如实说明「数据不足，无法判断趋势」\n"
            "- 所有数字结论必须来自上方数据或工具返回的真实数据，严禁编造\n"
            "- 回答用中文，结论放前面：先说进步还是退步、幅度多少，再给前后数据对比，最后给建议"
        ),
        "tools": ["query_student_trend", "query_student"],
        "max_tool_rounds": 4,
        "data_prep": _prep_trend,
    },
    "knowledge_gap": {
        "system_prompt": (
            "你是教学薄弱知识点分析助手，面向教师。\n"
            "上方提供了班级高频错误/薄弱知识点与 MOOC 知识库命中；工具 query_knowledge_gap 查薄弱点、"
            "search_knowledge_base 查知识点详情。\n"
            "规则：\n"
            "- 只基于上方数据与工具结果回答，没有数据支撑的知识点规律不要臆断\n"
            "- 把错题/薄弱点按知识点归类，指出高频项，给出补救建议（建议要有可操作性）\n"
            "- 回答用中文，结论放前面"
        ),
        "tools": ["search_knowledge_base", "query_knowledge_gap"],
        "max_tool_rounds": 4,
        "data_prep": _prep_knowledge_gap,
    },
    "warning": {
        "system_prompt": (
            "你是学生成绩预警分析助手，面向教师。\n"
            "上方提供了薄弱学生 Top 与已有预警预测结论；工具 query_student(weak_only=true) 查薄弱学生、"
            "query_knowledge_gap 查薄弱知识点。\n"
            "规则：\n"
            "- 只有数据支持时才下预警结论（薄弱率≥30%、平均分低、已有预测结论等）；数据不足以预警就如实说明\n"
            "- 对每个预警学生给出风险原因（哪些指标异常）和可操作建议\n"
            "- 回答用中文，按风险从高到低列出"
        ),
        "tools": ["query_student", "query_knowledge_gap"],
        "max_tool_rounds": 5,
        "data_prep": _prep_warning,
    },
    "attendance": {
        "system_prompt": (
            "你是课堂考勤分析助手，面向教师。\n"
            "上方提供了最近上传的考勤数据；工具 query_attendance 可补充查询。没有可用考勤数据时，"
            "请提示老师先上传考勤文件。\n"
            "规则：\n"
            "- 所有结论必须来自上方考勤数据或工具返回，严禁编造缺勤/迟到名单\n"
            "- 回答用中文，列出缺勤/迟到情况，需要时给出重点关注建议"
        ),
        "tools": ["query_attendance"],
        "max_tool_rounds": 3,
        "data_prep": _prep_attendance,
    },
    "report": {
        "system_prompt": (
            "你是教学报告写作助手，面向教师。\n"
            "你有全部查询工具（query_student / query_class_aggregate / query_student_trend / "
            "query_knowledge_gap / query_attendance / search_knowledge_base）。\n"
            "撰写前请先规划：报告需要哪些数据（班级整体、成绩分布、薄弱环节、趋势、考勤等），"
            "然后逐步调用工具获取真实数据，最后组织成完整报告。\n"
            "要求：\n"
            "- 报告用清晰的 Markdown 结构：一级标题为报告名称，下设 班级整体情况 / 成绩数据分析 / "
            "薄弱环节与风险 / 建议措施 等章节，数据用表格呈现\n"
            "- 报告内容必须全部来自工具返回或上方数据，严禁编造数字；数据缺失的章节如实说明\n"
            "- 班级名若是合并班（如「软件1-3班」=含1/2/3班）需在报告中说明粒度\n"
            "- 报告完成后，系统会自动将 Markdown 转成 Word 文档供老师下载，请保证结构完整、可直接成文"
        ),
        "tools": ["query_student", "query_class_aggregate", "query_student_trend",
                  "query_knowledge_gap", "query_attendance", "search_knowledge_base"],
        "max_tool_rounds": 6,
        "data_prep": _prep_report,
    },
}


# ============================================================
# 报告类：Markdown → Word（python-docx）
# ============================================================
def _strip_inline(s: str) -> str:
    """去掉行内 Markdown 标记（加粗/斜体/行内代码/链接），保留纯文本。"""
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"\*(.+?)\*", r"\1", s)
    s = re.sub(r"`(.+?)`", r"\1", s)
    s = re.sub(r"\[(.+?)\]\([^)]*\)", r"\1", s)
    return s.strip()


def _add_table(doc, block: List[List[str]]) -> None:
    rows = [r for r in block if not (len(r) == 1 and re.fullmatch(r":?-{3,}:?", r[0] or ""))]
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    try:
        from docx import Document
    except Exception:  # pragma: no cover - 依赖缺失时降级
        logger.exception("python-docx 未安装，无法生成 Word")
        return
    table = doc.add_table(rows=0, cols=ncols)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for r in rows:
        cells = table.add_row().cells
        for j in range(ncols):
            cells[j].text = _strip_inline(r[j] if j < len(r) else "")


def markdown_to_docx(md_text: str, out_path: Path) -> None:
    """把报告 Markdown 转成 Word：标题 / 表格 / 列表 / 段落。结构按内容自适应。"""
    from docx import Document

    doc = Document()
    lines = md_text.splitlines()
    i, n = 0, len(lines)
    while i < n:
        line = lines[i].rstrip().strip()
        if not line:
            i += 1
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(_strip_inline(m.group(2)), level=level)
            i += 1
            continue
        if line.startswith("|"):
            block: List[List[str]] = []
            while i < n and lines[i].strip().startswith("|"):
                block.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            _add_table(doc, block)
            continue
        m = re.match(r"^([-*+]|\d+[.、])\s+(.*)$", line)
        if m:
            doc.add_paragraph(_strip_inline(m.group(2)), style="List Bullet")
            i += 1
            continue
        doc.add_paragraph(_strip_inline(line))
        i += 1
    doc.save(str(out_path))


def generate_report_docx(md_text: str, title: str = "教学报告") -> Dict[str, str]:
    """
    把 AI 生成的报告 Markdown 转成 Word 存到 config.report_dir。
    返回 {"filename", "url_path"}；失败返回 {}（调用方不抛错，仅降级为无下载）。
    """
    if not md_text or not md_text.strip():
        return {}
    safe_title = re.sub(r'[\\/:*?"<>|\r\n\t ]+', "_", title)[:40]
    filename = f"{datetime.now():%Y%m%d%H%M%S}_{safe_title or '教学报告'}.docx"
    out = config.ai_report_dir / filename
    try:
        markdown_to_docx(md_text, out)
    except Exception as e:
        logger.exception("报告 Word 生成失败: %s", e)
        return {}
    logger.info("报告 Word 已生成：{}", out)
    return {"filename": filename, "url_path": f"/api/chat/report-download/{filename}"}
